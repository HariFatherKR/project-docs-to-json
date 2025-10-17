from __future__ import annotations

import logging
from io import BytesIO
from typing import List

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    docx = None  # type: ignore

from .. import settings
from ..utils import ocr

logger = logging.getLogger(__name__)


def extract_docx_blocks(data: bytes) -> tuple[str | None, List[str], List[List[str]]]:
    """Extract title, paragraphs, and tables from DOCX payload."""
    if docx is None:
        raise RuntimeError(
            "python-docx is not installed. Install dependencies from requirements.txt."
        )

    document = docx.Document(BytesIO(data))

    title: str | None = None
    paragraphs: List[str] = []
    table_rows: List[List[str]] = []

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if index == 0:
            title = text
        else:
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            table_rows.append([cell.text.strip() for cell in row.cells])

    if settings.ENABLE_OCR and ocr.available():
        paragraphs.extend(_extract_inline_image_text(document))

    return title, paragraphs, table_rows


def _extract_inline_image_text(document) -> List[str]:
    texts: List[str] = []

    for inline_shape in getattr(document, "inline_shapes", []):
        try:
            blip = inline_shape._inline.graphic.graphicData.pic.blipFill.blip  # type: ignore[attr-defined]
            embed = getattr(blip, "embed", None)
            link = getattr(blip, "link", None)
            rel_id = embed or link
            if not rel_id:
                continue
            image_part = document.part.related_parts.get(rel_id)
            if image_part is None:
                continue
            text = ocr.bytes_to_text(image_part.blob)
            if text:
                texts.append(text)
        except Exception as exc:  # pragma: no cover
            logger.warning("docx-ocr-inline-shape-failed", exc_info=exc)
            continue

    return texts
